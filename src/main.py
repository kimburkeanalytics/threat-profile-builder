from datetime import date
from pathlib import Path
import re

from dotenv import load_dotenv
from openai import OpenAI

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

# --------------------
# Setup
# --------------------
load_dotenv()
client = OpenAI()

actor = input("Which threat actor folder should I use? ").strip().lower()

actor_display_names = {
    "cl0p": "Cl0p Ransomware",
    "lockbit": "LockBit Ransomware",
    "lazarus": "Lazarus Group",
    "scattered_spider": "Scattered Spider",
}

actor_display = actor_display_names.get(actor, actor.replace("_", " ").title())

data_folder = Path(f"data/{actor}")
output_folder = Path("output")
output_folder.mkdir(exist_ok=True)

master_file = output_folder / f"{actor}_master_source.txt"
brief_file = output_folder / f"{actor}_ai_generated_brief.md"
docx_file = output_folder / f"{actor}_threat_profile.docx"

# --------------------
# Validate
# --------------------
if not data_folder.exists():
    raise FileNotFoundError(f"Missing folder: {data_folder}")

source_files = list(data_folder.glob("*.txt"))

if not source_files:
    raise FileNotFoundError(f"No .txt files found in {data_folder}")

# --------------------
# Build sourced corpus
# --------------------
combined = []
source_map = {}

for i, file in enumerate(source_files, start=1):
    source_id = f"S{i}"
    source_map[source_id] = file.name
    text = file.read_text(encoding="utf-8")

    combined.append(f"[{source_id}] SOURCE FILE: {file.name}")
    combined.append(text)
    combined.append("\n")

source_text = "\n".join(combined)
master_file.write_text(source_text, encoding="utf-8")

# --------------------
# Generate AI brief with source markers
# --------------------
prompt = f"""
You are a senior cyber threat intelligence analyst.

Using ONLY the source material below, write a professional cyber threat profile on {actor_display}.

Use clean markdown headings only: # and ##.

Report title must be:
# Cyber Threat Profile: {actor_display}

Use exactly these sections, in this order:

## Executive Summary
## Impact
## Origin Story
## Operational Model
## Capabilities and TTPs
## Key Partners & Affiliations
## Vulnerabilities
## Outlook

Section guidance:
- - Executive Summary: true summary only. Do not introduce facts that are not reflected in the sections below. Write exactly one paragraph, 2-4 sentences maximum. No bullets.
- Impact: explain why this actor matters. Address sectors/industries targeted, major attacks, damage, business impact, victim scale, or strategic significance when supported by the sources.
- Origin Story: explain when the actor first emerged, potential predecessor groups, suspected links to earlier clusters, and WHY analysts assess those links. Do not simply proclaim connections without explaining the evidence.
- Operational Model: focus on structure and business model. Address whether the actor appears closed, affiliate-based, ransomware-as-a-service, scalable, opportunistic, centralized, or specialized, when supported by sources.
- Capabilities and TTPs: explain technical and operational capabilities. Include a relative capability ranking compared with other ransomware/extortion groups when the source material supports it, and clearly state the basis for that assessment.
- Key Partners & Affiliations: include known or suspected links to other groups, criminal networks, ransomware families, affiliates, infrastructure providers, money laundering ecosystems, or foreign governments/state-linked activity when supported by sources.
- Vulnerabilities: identify operational, technical, organizational, reputational, law-enforcement, infrastructure, monetization, or strategic weaknesses when supported. If sources do not support meaningful vulnerabilities, say so.
- Outlook: assess where the actor is trending.

Requirements:
- Use only the supplied source material.
- Do not invent facts.
- Add source markers like [S1], [S2], or [S3] at the end of important factual paragraphs.
- If a claim is based on multiple sources, cite multiple markers like [S1][S3].
- Do not use heading levels deeper than ##.
- Keep tone executive-ready and analytical.

SOURCE MATERIAL:
{source_text}
"""

response = client.chat.completions.create(
    model="gpt-4o-mini",
    messages=[{"role": "user", "content": prompt}],
)

brief = response.choices[0].message.content
brief_file.write_text(brief, encoding="utf-8")

# --------------------
# Helpers
# --------------------
def add_markdown_text(paragraph, text):
    """Convert **bold** markdown into real bold text."""
    parts = re.split(r"(\*\*.*?\*\*)", text)

    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def clean_heading(line):
    return line.lstrip("#").strip().replace("**", "")


def split_exec_summary(markdown_text):
    """Separate Executive Summary content from the rest of the markdown report."""
    lines = markdown_text.splitlines()
    exec_lines = []
    body_lines = []
    in_exec = False

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("# Cyber Threat Profile"):
            continue

        if stripped == "## Executive Summary":
            in_exec = True
            continue

        if stripped.startswith("## ") and in_exec:
            in_exec = False
            body_lines.append(line)
            continue

        if in_exec:
            exec_lines.append(line)
        else:
            body_lines.append(line)

    return "\n".join(exec_lines).strip(), "\n".join(body_lines).strip()


exec_summary, body_text = split_exec_summary(brief)

# --------------------
# Build Word document
# --------------------
doc = Document()

section = doc.sections[0]
section.top_margin = Inches(0.85)
section.bottom_margin = Inches(0.85)
section.left_margin = Inches(0.9)
section.right_margin = Inches(0.9)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(11)

# --------------------
# Cover page
# --------------------
for _ in range(9):
    doc.add_paragraph("")

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run(f"Cyber Threat Profile:\n{actor_display}")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(31, 78, 121)

for _ in range(2):
    doc.add_paragraph("")

brand = doc.add_paragraph()
brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = brand.add_run("THREAT PROFILE BUILDER")
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(90, 90, 90)

generated = doc.add_paragraph()
generated.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = generated.add_run(f"Generated {date.today().isoformat()}")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(90, 90, 90)

doc.add_page_break()

# --------------------
# Executive Summary Box
# --------------------
summary_table = doc.add_table(rows=1, cols=1)
summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
summary_cell = summary_table.cell(0, 0)
summary_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

heading_p = summary_cell.paragraphs[0]
heading_run = heading_p.add_run("Executive Summary")
heading_run.bold = True
heading_run.font.size = Pt(14)
heading_run.font.color.rgb = RGBColor(31, 78, 121)

for line in exec_summary.splitlines():
    line = line.strip()
    if not line:
        continue

    if line.startswith("- "):
        p = summary_cell.add_paragraph(style="List Bullet")
        add_markdown_text(p, line[2:])
    else:
        p = summary_cell.add_paragraph()
        add_markdown_text(p, line.replace("## Executive Summary", ""))

    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0

doc.add_paragraph("")

# --------------------
# Body
# --------------------
for line in body_text.splitlines():
    line = line.strip()

    if not line:
        continue

    if line.startswith("#"):
        heading = clean_heading(line)

        if heading.lower().startswith("cyber threat profile"):
            continue

        p = doc.add_heading(heading, level=1)
        p.runs[0].font.color.rgb = RGBColor(31, 78, 121)
        p.runs[0].font.size = Pt(15)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)

    elif line.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        add_markdown_text(p, line[2:])
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0

    else:
        p = doc.add_paragraph()
        add_markdown_text(p, line)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.0

# --------------------
# Source Appendix
# --------------------
doc.add_page_break()

p = doc.add_heading("Appendix A: Source Key", level=1)
p.runs[0].font.color.rgb = RGBColor(31, 78, 121)

for source_id, filename in source_map.items():
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"[{source_id}] ").bold = True
    p.add_run(filename)
    p.paragraph_format.space_after = Pt(2)

# --------------------
# Footer
# --------------------
for sec in doc.sections:
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = "Threat Profile Builder | LLM-assisted cyber threat intelligence workflow"

doc.save(docx_file)

print("Successfully generated:")
print(f"- {master_file}")
print(f"- {brief_file}")
print(f"- {docx_file}")